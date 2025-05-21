import tkinter as tk
from tkinter import scrolledtext, messagebox, filedialog, simpledialog, ttk
import numpy as np
import pandas as pd
import pytesseract
from PIL import Image
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.cluster import SpectralClustering
import os
import docx
import re
from collections import defaultdict
import unicodedata
from sklearn.metrics import silhouette_score


# Cấu hình Tesseract OCR
pytesseract.pytesseract.tesseract_cmd = r"C:\\Program Files\\Tesseract-OCR\\tesseract.exe"

# Màu sắc giao diện
BACKGROUND_COLOR = "#D0F0C0"
BUTTON_COLOR = "#4CAF50"
BUTTON_HOVER_COLOR = "#45a049"
TEXT_COLOR = "#333333"
HEADER_COLOR = "#5f6368"
PROGRESS_COLOR = "#FF9800"

def clean_text(text):
    text = text.lower()
    text = unicodedata.normalize('NFC', text)
    text = re.sub(r'[^\w\s]', '', text, flags=re.UNICODE)
    text = re.sub(r'\s+', ' ', text).strip()
    return text

def format_table(cells):
    col_widths = [max(len(cell) for cell in column) for column in zip(*cells)]
    rows = []
    for row in cells:
        formatted_row = " | ".join(f"{cell:<{col_widths[i]}}" for i, cell in enumerate(row))
        rows.append(f"| {formatted_row} |")
    separator = "|" + "|".join("-" * (w + 2) for w in col_widths) + "|"
    return "\n".join([rows[0], separator] + rows[1:])

# Danh sách chứa tên tài liệu tương ứng với từng dòng
doc_sources = []

from sklearn.metrics import silhouette_score
from sklearn.cluster import SpectralClustering

def cluster_documents():
    raw_lines = text_input.get("1.0", tk.END).strip().split("\n")
    docs = []
    is_table = []
    tables = []
    doc_sources = []  # Đảm bảo doc_sources được khởi tạo lại cho mỗi lần phân loại

    # Gán nguồn tài liệu cho mỗi dòng
    current_doc_name = ""
    parsed_lines = []
    for line in raw_lines:
        if line.startswith("--- [TÀI LIỆU: "):
            match = re.match(r"--- \[TÀI LIỆU: (.+?)\] ---", line)
            if match:
                current_doc_name = match.group(1).strip()
        elif line.strip():
            parsed_lines.append((line.strip(), current_doc_name))

    for line, doc_name in parsed_lines:
        if not doc_name:  # Kiểm tra xem nếu dòng không có tên tài liệu
            doc_name = "Không rõ nguồn"  # Gán tên mặc định nếu không có tên tài liệu

        if "\t" in line:
            cells = [cell.strip() for cell in line.split("\t")]
            docs.append(' '.join(cells))
            is_table.append(True)
            tables.append(cells)
        else:
            docs.append(clean_text(line))
            is_table.append(False)
            tables.append(None)
        doc_sources.append(doc_name)

    if not docs:
        messagebox.showerror("Lỗi", "Vui lòng nhập dữ liệu!")
        return

    # Tính số cụm tối ưu bằng cách sử dụng Silhouette Score
    vectorizer = TfidfVectorizer(stop_words='english')
    X_tfidf = vectorizer.fit_transform(docs)
    
    best_n_clusters = 2  # Khởi tạo với 2 cụm
    best_score = -1  # Score thấp nhất có thể

    for n_clusters in range(2, min(50, len(docs))):  # Kiểm tra số cụm từ 2 đến 50
        spectral = SpectralClustering(n_clusters=n_clusters, affinity='precomputed', assign_labels='kmeans')
        similarity_matrix = cosine_similarity(X_tfidf)
        labels = spectral.fit_predict(similarity_matrix)
        
        # Tính Silhouette Score
        score = silhouette_score(X_tfidf, labels)
        if score > best_score:
            best_score = score
            best_n_clusters = n_clusters

    # Phân cụm với số cụm tối ưu
    spectral = SpectralClustering(n_clusters=best_n_clusters, affinity='precomputed', assign_labels='kmeans')
    similarity_matrix = cosine_similarity(X_tfidf)
    labels = spectral.fit_predict(similarity_matrix)

    cluster_dict = defaultdict(list)
    for i, (original_line, doc_name) in enumerate(parsed_lines):
        cluster_dict[labels[i]].append((original_line, is_table[i], tables[i], doc_sources[i]))

    output_text.delete("1.0", tk.END)
    result = ""

    for cluster, items in sorted(cluster_dict.items()):
        result += f"\n\n🔹 Cụm {cluster}:\n"
        result += "-" * 60 + "\n"
        for text, table_flag, cells, doc_name in items:
            if table_flag and cells:
                result += f"[{doc_name}]\n"
                result += format_table([cells]) + "\n"
            else:
                result += f"[{doc_name}] - {text}\n"
        result += "-" * 60 + "\n"

    output_text.insert(tk.END, result.strip())
    progress_bar.stop()
    messagebox.showinfo("Hoàn tất", f"Phân loại hoàn tất với {best_n_clusters} cụm!")


def load_multiple_files():
    global doc_sources
    file_paths = filedialog.askopenfilenames(filetypes=[("Supported Files", "*.txt;*.docx;*.csv;*.png;*.jpg;*.jpeg")])
    if not file_paths:
        return

    all_text = []
    doc_sources = []

    for file_path in file_paths:
        file_extension = os.path.splitext(file_path)[-1].lower()
        try:
            text = ""
            file_name = os.path.basename(file_path)
            if file_extension == ".txt":
                with open(file_path, "r", encoding="utf-8") as file:
                    lines = file.read().splitlines()
            elif file_extension == ".docx":
                doc = docx.Document(file_path)
                lines = [clean_text(para.text.strip()) for para in doc.paragraphs if para.text.strip()]
                for table in doc.tables:
                    for row in table.rows:
                        row_text = '\t'.join(clean_text(cell.text.strip()) for cell in row.cells)
                        if row_text.strip():
                            lines.append(row_text)
            elif file_extension == ".csv":
                df = pd.read_csv(file_path)
                lines = df.astype(str).apply(lambda row: "\t".join(row.values), axis=1).tolist()
            elif file_extension in [".png", ".jpg", ".jpeg"]:
                image = Image.open(file_path)
                text = pytesseract.image_to_string(image, lang='vie')
                lines = text.splitlines()
            else:
                continue

            if lines:
                all_text.append(f"--- [TÀI LIỆU: {file_name}] ---")
                all_text.extend(lines)
        except Exception as e:
            messagebox.showerror("Lỗi", f"Lỗi đọc {file_path}: {str(e)}")

    text_input.delete("1.0", tk.END)
    text_input.insert(tk.END, "\n".join(all_text))

def clear_text():
    text_input.delete("1.0", tk.END)
    output_text.delete("1.0", tk.END)
    text = re.sub(r'^[\ - \•\*\u2022]+[\s]*', '', text) 

def load_from_file():
    file_path = filedialog.askopenfilename(filetypes=[("Text Files", "*.txt"), ("Word Documents", "*.docx"), ("CSV Files", "*.csv"), ("Image Files", "*.png;*.jpg;*.jpeg")])
    if not file_path:
        return

    file_extension = os.path.splitext(file_path)[-1].lower()
    try:
        text = ""
        file_name = os.path.basename(file_path)
        if file_extension == ".txt":
            with open(file_path, "r", encoding="utf-8") as file:
                lines = file.read().splitlines()
        elif file_extension == ".docx":
            doc = docx.Document(file_path)
            lines = [clean_text(para.text.strip()) for para in doc.paragraphs if para.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = '\t'.join(clean_text(cell.text.strip()) for cell in row.cells)
                    if row_text.strip():
                        lines.append(row_text)
        elif file_extension == ".csv":
            df = pd.read_csv(file_path)
            lines = df.astype(str).apply(lambda row: "\t".join(row.values), axis=1).tolist()
        elif file_extension in [".png", ".jpg", ".jpeg"]:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image, lang='vie')
            lines = text.splitlines()
        else:
            messagebox.showerror("Lỗi", "Định dạng tệp không được hỗ trợ!")
            return

        if lines:
            text_input.delete("1.0", tk.END)
            text_input.insert(tk.END, f"--- [TÀI LIỆU: {file_name}] ---\n" + "\n".join(lines))
    except Exception as e:
        messagebox.showerror("Lỗi", f"Không thể đọc tệp: {str(e)}")
def export_results():
    content = output_text.get("1.0", tk.END).strip()
    if not content:
        messagebox.showwarning("Thông báo", "Không có dữ liệu để lưu!")
        return

    file_path = filedialog.asksaveasfilename(defaultextension=".txt",
                                             filetypes=[("Text File", "*.txt"), ("Word Document", "*.docx")])
    if not file_path:
        return

    try:
        if file_path.endswith(".txt"):
            with open(file_path, "w", encoding="utf-8") as f:
                f.write(content)
        elif file_path.endswith(".docx"):
            doc = docx.Document()
            for line in content.split("\n"):
                doc.add_paragraph(line)
            doc.save(file_path)
        messagebox.showinfo("Thành công", "Đã lưu kết quả thành công!")
    except Exception as e:
        messagebox.showerror("Lỗi", f"Không thể lưu kết quả: {str(e)}")


# Giao diện người dùng
root = tk.Tk()
root.title("Phân loại tài liệu - Spectral Clustering")
root.geometry("1000x600")
root.state("zoomed")
root.configure(bg=BACKGROUND_COLOR)

style = ttk.Style()
style.configure("TButton", background=BUTTON_COLOR, foreground="black", font=("Arial", 12))
style.map("TButton", background=[('active', BUTTON_HOVER_COLOR)])
style.configure("TProgressbar", thickness=20, troughcolor=BACKGROUND_COLOR, background=PROGRESS_COLOR)
style.configure("Custom.TFrame", background=BACKGROUND_COLOR)

notebook = ttk.Notebook(root)
notebook.pack(fill=tk.BOTH, expand=True)

data_frame = ttk.Frame(notebook, style="Custom.TFrame")
notebook.add(data_frame, text="📂 Tải dữ liệu")

load_button = ttk.Button(data_frame, text="📥 Chọn file", command=load_from_file)
load_button.pack(pady=10)

text_input = scrolledtext.ScrolledText(data_frame, wrap=tk.WORD, height=20, bg="#FFFFFF", fg=TEXT_COLOR, font=("Arial", 12))
text_input.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

task_frame = ttk.Frame(notebook, style="Custom.TFrame")
notebook.add(task_frame, text="🔍 Phân loại")

cluster_button = ttk.Button(task_frame, text="🚀 Phân loại", command=cluster_documents)
cluster_button.pack(pady=10)

output_text = scrolledtext.ScrolledText(task_frame, wrap=tk.WORD, height=20, bg="#FFFFFF", fg=TEXT_COLOR, font=("Arial", 12))
output_text.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

load_files_button = ttk.Button(task_frame, text="📂 Tải nhiều file", command=load_multiple_files)
load_files_button.pack(pady=5)

save_button = ttk.Button(task_frame, text="💾 Lưu kết quả", command=export_results)
save_button.pack(pady=5)


clear_button = ttk.Button(task_frame, text="🗑 Xóa", command=clear_text)
clear_button.pack(pady=5)

progress_bar = ttk.Progressbar(task_frame, mode='indeterminate', style="TProgressbar")
progress_bar.pack(fill=tk.X, padx=10, pady=5)

root.mainloop()
